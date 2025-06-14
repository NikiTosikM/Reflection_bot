import json
from datetime import date

import yadisk
from yadisk import Client
from docx import Document


class FilePhrases:
    @staticmethod
    def read_file(file_name):
        with open(file_name, encoding='utf-8') as file:
            file = json.load(file)
            return file
        
    @classmethod
    def stages_ref_text(cls, file_name, name_stage):
        file = cls.read_file(file_name)
        answer_text = file['stages_writing_reflection'][name_stage]
        return answer_text
    
    @classmethod
    def reflection_text(cls, file_name):
        file = cls.read_file(file_name)
        return file['write_reflection']
        
    @classmethod
    def get_phrase_when_start_bot(cls, file_name):
        file = cls.read_file(file_name)
        return file['start']
    
    @classmethod
    def display_beautiful_format(cls, file_name):
        def get_block():
            for block in cls.read_file(file_name)["reflection_phrases"]:
                yield block
        result_phrases_format = ''
        for block in get_block():
            for phrase in block:
                result_phrases_format += phrase + '\n'
            result_phrases_format += '\n\n'
        
        return result_phrases_format


class WordDocument():
    object_ = None
    
    def __new__(cls):
        if cls.object_ is None:
            cls.object_= super().__new__(cls)
            return cls.object_
        else:
            return cls.object_
            
    def __init__(self):
        self.doc_obj = Document()
    
    def add_text_document(self, text, name_user):
        if not self.doc_obj.paragraphs:
            date_ = date.today()
            self.doc_obj.add_heading(f'{date_}', level=1)
        self.doc_obj.add_heading(f'{name_user}', level=3)
        self.doc_obj.add_paragraph(text+'\n\n')
        
    def save_file(self):
        file_name = f'{date.today()}.docx'
        self.doc_obj.save(file_name)
        
        return file_name
        
        
class YaDisk():
    def __init__(self, token, doc_obj):
        self.token = token
        self.client: Client = yadisk.AsyncClient(token=self.token)
        self.doc: WordDocument = doc_obj
        
    async def send_docx_document(self):
        file_name = self.doc.save_file()
        async with self.client:
            await self.client.upload(
                file_name, 
                f'Reflection/{file_name}'
            )
        for paragraph in list(self.doc.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
            
            
        
    